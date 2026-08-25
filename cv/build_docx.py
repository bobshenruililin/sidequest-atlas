#!/usr/bin/env python3
"""Build the DOCX twin of shen-ruililin-cv.tex. Content must stay in lockstep."""

from __future__ import annotations

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor

OUT = "shen-ruililin-cv.docx"
PAGE_WIDTH_MM = 210
MARGIN_MM = 11.5


def set_run_font(run, name="Times New Roman", size=10, bold=False, italic=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor(0, 0, 0)


def add_bottom_border(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "11")
    bottom.set(qn("w:space"), "3")
    bottom.set(qn("w:color"), "000000")
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_tab_right(paragraph, pos_twips):
    pPr = paragraph._p.get_or_add_pPr()
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "right")
    tab.set(qn("w:pos"), str(pos_twips))
    tabs.append(tab)
    pPr.append(tabs)


def usable_width_twips(section):
    return int(section.page_width.pt * 20) - int(section.left_margin.pt * 20) - int(
        section.right_margin.pt * 20
    )


def heading(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(4)
    pf.space_after = Pt(0.5)
    pf.line_spacing = Pt(11)
    run = p.add_run(text.upper())
    set_run_font(run, size=11, bold=True)
    add_bottom_border(p)
    return p


def head_row(doc, left, right, left2, right2, width_twips):
    p1 = doc.add_paragraph()
    pf = p1.paragraph_format
    pf.space_before = Pt(1.6)
    pf.space_after = Pt(0)
    pf.line_spacing = Pt(11.5)
    add_tab_right(p1, width_twips)
    r = p1.add_run(left)
    set_run_font(r, size=10.5, bold=True)
    p1.add_run("\t")
    r2 = p1.add_run(right)
    set_run_font(r2, size=10.5, bold=True)

    p2 = doc.add_paragraph()
    pf2 = p2.paragraph_format
    pf2.space_before = Pt(0)
    pf2.space_after = Pt(0.2)
    pf2.line_spacing = Pt(11.5)
    add_tab_right(p2, width_twips)
    r3 = p2.add_run(left2)
    set_run_font(r3, size=10)
    p2.add_run("\t")
    r4 = p2.add_run(right2)
    set_run_font(r4, size=10)
    return p1, p2


def add_mixed(paragraph, parts, size=10):
    for text, bold in parts:
        run = paragraph.add_run(text)
        set_run_font(run, size=size, bold=bold)
    return paragraph


def bullet(doc, parts):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0.15)
    pf.line_spacing = Pt(11.5)
    pf.left_indent = Mm(4.5)
    pf.first_line_indent = Mm(-3.2)
    add_mixed(p, [("•  ", False), *parts], size=10)
    return p


def honor(doc, parts):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0.5)
    pf.line_spacing = Pt(11.5)
    add_mixed(p, parts, size=10)
    return p


def main():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Mm(PAGE_WIDTH_MM)
    section.page_height = Mm(297)
    section.left_margin = Mm(MARGIN_MM)
    section.right_margin = Mm(MARGIN_MM)
    section.top_margin = Mm(9.5)
    section.bottom_margin = Mm(9.5)
    width = usable_width_twips(section)

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(10)

    name = doc.add_paragraph()
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name.paragraph_format.space_before = Pt(0)
    name.paragraph_format.space_after = Pt(2)
    name.paragraph_format.line_spacing = Pt(15)
    nr = name.add_run("Shen Ruililin")
    set_run_font(nr, size=16, bold=True)

    ident = doc.add_paragraph()
    ident.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ident.paragraph_format.space_before = Pt(0)
    ident.paragraph_format.space_after = Pt(2)
    ident.paragraph_format.line_spacing = Pt(12)
    ir = ident.add_run(
        "BASc Global Health and Development, HKU  ·  Year 2  ·  CGPA 4.09/4.30"
    )
    set_run_font(ir, size=10, bold=True)

    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.paragraph_format.space_before = Pt(0)
    contact.paragraph_format.space_after = Pt(2)
    contact.paragraph_format.line_spacing = Pt(12)
    cr = contact.add_run(
        "bobshenruililin@gmail.com  ·  shenrll@connect.hku.hk  ·  "
        "+852 5573 6530  ·  linkedin.com/in/shenruililin"
    )
    set_run_font(cr, size=10)

    heading(doc, "Education")
    head_row(
        doc,
        "The University of Hong Kong",
        "Hong Kong",
        "BASc Global Health and Development, HKUMed · Year 2",
        "2025–2029",
        width,
    )
    bullet(
        doc,
        [
            ("CGPA ", False),
            ("4.09/4.30", True),
            (" (current). ", False),
            ("St. John's College", True),
            (" (2025–). Public health, quantitative methods, and development.", False),
        ],
    )
    head_row(
        doc,
        "Po Leung Kuk Choi Kai Yau School",
        "Hong Kong",
        "IBDP 41/45 · IGCSE 8A*, 2A",
        "2019–2025",
        width,
    )
    bullet(
        doc,
        [
            ("Academic Scholarships", True),
            (
                " (2022–23, 2024–25). Co-founder, Medical Club (2023–24; membership to 80+).",
                False,
            ),
        ],
    )

    heading(doc, "Honours and Awards")
    honor(
        doc,
        [
            ("Tung and Ngai Foundation Scholarship", True),
            (
                " (2025–; 1 of 4 HKU scholars). Full HKU tuition; living ",
                False,
            ),
            ("HK$40,000/year", True),
            ("; overseas/Mainland subsidy ", False),
            ("max HK$80,000", True),
            ("; Endeavour Support Fund ", False),
            ("up to HK$100,000", True),
            (".", False),
        ],
    )
    honor(
        doc,
        [
            ("Laidlaw Undergraduate Research and Leadership Programme", True),
            (
                ", 2025–26. Heat × cardiovascular admissions among older adults; Prof. ",
                False,
            ),
            ("David Bishai", True),
            (", HKU School of Public Health.", False),
        ],
    )
    honor(
        doc,
        [
            ("Martin Scholar", True),
            (", St. John's College, The University of Hong Kong, 2025–26.", False),
        ],
    )
    honor(
        doc,
        [
            ("Tam Wun Tsun HKU Horizons Student Enrichment Award", True),
            (", 2025–26.", False),
        ],
    )
    honor(
        doc,
        [
            ("MIT Hong Kong Innovation Node Youth Fellowship", True),
            (", 2023, Hong Kong.", False),
        ],
    )
    honor(
        doc,
        [
            ("IYPT", True),
            (
                " (Hong Kong): Champion 2022 (1st/16); first runner-up 2021 (2nd/8).",
                False,
            ),
        ],
    )
    honor(
        doc,
        [
            ("Hong Kong Young Writers' Award", True),
            (", Bauhinia Club Award, 2024 (1st/1400+).", False),
        ],
    )

    heading(doc, "Research Experience")
    head_row(
        doc,
        "Laidlaw Scholars Programme, HKU",
        "Hong Kong",
        "Undergraduate research scholar",
        "Jan 2026–present",
        width,
    )
    bullet(
        doc,
        [
            (
                "Heat and cardiovascular admissions among older adults using ",
                False,
            ),
            ("Hospital Authority", True),
            (
                " data, with Prof. David Bishai, HKU School of Public Health. Ongoing analysis toward dual-season hospital preparedness and cross-boundary heat-health surveillance.",
                False,
            ),
        ],
    )
    bullet(
        doc,
        [
            ("Forthcoming, Oct 2026: ", False),
            ("invited speaker", True),
            (
                ", Global Conference on Environmental Science and Technology (GEST), Valencia; ",
                False,
            ),
            ("attending", True),
            (" the Laidlaw Scholars Annual Conference, London.", False),
        ],
    )
    head_row(
        doc,
        "Hong Kong Observatory and PolyU MicroLARGE Lab",
        "Hong Kong",
        "Research intern, Junior Researcher Mentoring Programme",
        "Feb–Jul 2024",
        width,
    )
    bullet(
        doc,
        [
            (
                "Modelled GraphCast integration into Hong Kong forecasting; atmospheric-rivers identification presented at AGU.",
                False,
            )
        ],
    )
    head_row(
        doc,
        "HKUST IAS Center for Quantum Technologies",
        "Hong Kong",
        "Research intern, Quantum Computing for Gifted Students Scheme",
        "Jul 2023–Mar 2024",
        width,
    )
    bullet(
        doc,
        [
            (
                "Qiskit models and SpinQ auto-calibration; presented entanglement versus measurement work at physics-department seminars.",
                False,
            )
        ],
    )
    head_row(
        doc,
        "MIT Hong Kong Innovation Node",
        "Hong Kong",
        "Youth Fellow",
        "Jul–Sep 2023",
        width,
    )
    bullet(
        doc,
        [
            (
                "Prototyped MEDocGPT, an LLM healthcare chatbot and mobile app for early health intervention.",
                False,
            )
        ],
    )

    heading(doc, "Leadership")
    head_row(
        doc,
        "Wu Zhi Qiao (WZQ, HKU)",
        "Macha, Gansu",
        "Student Coordinator (1 of 3)",
        "May–Jun 2026",
        width,
    )
    bullet(
        doc,
        [
            (
                "One of three student coordinators for on-site rural health and infrastructure service: logistics, listening, and turning village needs into feasible student work.",
                False,
            )
        ],
    )
    head_row(
        doc,
        "NEWDAY - Nansen East-West Dialogue Academy",
        "Lillehammer",
        "Selected participant, Nansen Academy (1 of 3 HKU)",
        "29 Jul–6 Aug 2026",
        width,
    )
    bullet(
        doc,
        [
            (
                "Nine-day residential East-West dialogue (East Asia x Nordics). Then independent Nordic fieldwork, Oslo-Stockholm-Helsinki, 6-20 Aug 2026 (not a second academy).",
                False,
            )
        ],
    )
    head_row(
        doc,
        "Hong Kong Laureate Forum",
        "Hong Kong",
        "Forum Ambassador",
        "Sep–Nov 2025",
        width,
    )
    bullet(
        doc,
        [
            (
                "Worked with 20+ Shaw Laureates, young scientists, and the HKLF Secretariat.",
                False,
            )
        ],
    )
    head_row(
        doc,
        "Asian Spring Program on Rationality",
        "Taoyuan",
        "Fully-funded scholar (1 of 20)",
        "Feb 2025",
        width,
    )
    bullet(
        doc,
        [
            (
                "Ten-day programme in game theory, Bayesian inference, and Fermi estimation.",
                False,
            )
        ],
    )

    heading(doc, "Skills")
    skills = doc.add_paragraph()
    skills.paragraph_format.space_before = Pt(0)
    skills.paragraph_format.space_after = Pt(0)
    skills.paragraph_format.line_spacing = Pt(12)
    add_mixed(
        skills,
        [
            ("Languages: ", True),
            ("Cantonese, Mandarin, English. ", False),
            ("Programming: ", True),
            ("Python, R, C++.", False),
        ],
        size=10,
    )

    doc.save(OUT)
    print(f"wrote {OUT}")


if __name__ == "__main__":
    main()
